#!/usr/bin/env python3
"""
parse_docx.py — Word .docx → classified.json для layout_designer.

Использует python-docx. Эвристики:
- Heading 1 → title slide (если первый) или divider
- Heading 2 → divider
- Heading 3 → content slide
- bullet/numbered list paragraphs → body
- tables → table-content
- inline images → image_path

Usage:
    python3 parse_docx.py <input.docx> <output_classified.json>
"""
import sys
import json
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def get_heading_level(para):
    """Возвращает 1/2/3 для Heading 1/2/3, иначе None."""
    style = para.style.name if para.style else ""
    if style.startswith("Heading 1"):
        return 1
    if style.startswith("Heading 2"):
        return 2
    if style.startswith("Heading 3"):
        return 3
    return None


def is_list_item(para):
    style = para.style.name if para.style else ""
    return ("List" in style) or ("Bullet" in style) or ("Number" in style)


def parse_docx(doc_path):
    doc = Document(doc_path)
    slides = []
    current = None
    body_buffer = []

    def flush():
        nonlocal current, body_buffer
        if current is not None:
            if body_buffer:
                current["body"] = body_buffer.copy()
            slides.append(current)
        current = None
        body_buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        h = get_heading_level(para)
        if h == 1:
            flush()
            cat = "title_open" if not slides else "divider"
            current = {"num": len(slides) + 1, "category": cat, "tone_hint": "light",
                       "title": text}
        elif h == 2:
            flush()
            current = {"num": len(slides) + 1, "category": "divider", "tone_hint": "dark",
                       "title": text, "number": f"{len(slides)+1:02d}"}
        elif h == 3:
            flush()
            current = {"num": len(slides) + 1, "category": "content_text", "tone_hint": "light",
                       "title": text}
        elif is_list_item(para):
            body_buffer.append(text)
        else:
            # Plain paragraph
            if current is None:
                current = {"num": len(slides) + 1, "category": "content_text", "tone_hint": "light",
                           "title": text[:60]}
            else:
                body_buffer.append(text)

    flush()

    # Auto-classify based on body length
    for s in slides:
        body = s.get("body", [])
        if s["category"] == "content_text" and len(body) >= 4:
            s["category"] = "content_4subtitles"
            s["subtitles"] = [b[:30] for b in body[:4]]
            s["col_bodies"] = body[:4]
        elif s["category"] == "content_text" and len(body) == 3:
            s["category"] = "content_3col"
            s["subtitles"] = [b[:30] for b in body[:3]]
            s["col_bodies"] = body[:3]

    # Tables (отдельно от paragraphs — добавим как отдельные slides в конец)
    for tbl in doc.tables:
        rows_data = []
        for row in tbl.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if rows_data:
            slides.append({
                "num": len(slides) + 1,
                "category": "table",
                "tone_hint": "light",
                "title": f"Таблица {len([s for s in slides if s.get('category') == 'table']) + 1}",
                "table_data": rows_data,
            })

    # Finale
    if slides and slides[-1]["category"] != "logo_finale":
        slides.append({"num": len(slides) + 1, "category": "logo_finale", "tone_hint": "green",
                       "title": "Спасибо!", "subtitle": "Cloud.ru"})

    return slides


def main():
    if len(sys.argv) < 3:
        print("Usage: parse_docx.py <input.docx> <output_classified.json>", file=sys.stderr)
        sys.exit(1)
    docx_path = sys.argv[1]
    out_path = sys.argv[2]

    slides = parse_docx(docx_path)
    classified = {"_source": docx_path, "_parser": "parse_docx.py", "slides": slides}

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(classified, f, ensure_ascii=False, indent=2)

    print(f"Parsed {docx_path}: {len(slides)} slides → {out_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
